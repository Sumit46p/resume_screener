"""
DOCX Parser with support for tables, headers, and complex layouts.
"""
import io
from typing import Optional
from docx import Document
from docx.table import Table


def parse_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text preserving structure
    """
    try:
        doc = Document(file_path)
        return _extract_text_from_document(doc)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_docx_from_bytes(content: bytes) -> str:
    """
    Extract text from DOCX bytes (for file uploads).
    
    Args:
        content: DOCX file content as bytes
        
    Returns:
        Extracted text preserving structure
    """
    try:
        doc = Document(io.BytesIO(content))
        return _extract_text_from_document(doc)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX from bytes: {str(e)}")


def _extract_text_from_document(doc: Document) -> str:
    """
    Extract text from document including paragraphs and tables.
    """
    text_parts = []
    
    # Process all body elements in order
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # It's a paragraph
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:
                        text_parts.append(text)
                    break
        elif element.tag.endswith('tbl'):
            # It's a table
            for table in doc.tables:
                if table._element == element:
                    table_text = _extract_table_text(table)
                    if table_text:
                        text_parts.append(table_text)
                    break
    
    # If the above iteration didn't work properly, fallback to simple extraction
    if not text_parts:
        text_parts = _simple_extract(doc)
    
    return '\n\n'.join(text_parts)


def _simple_extract(doc: Document) -> list:
    """
    Simple text extraction fallback.
    """
    text_parts = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)
    
    # Extract tables
    for table in doc.tables:
        table_text = _extract_table_text(table)
        if table_text:
            text_parts.append(table_text)
    
    return text_parts


def _extract_table_text(table: Table) -> str:
    """
    Extract text from a table, preserving row structure.
    """
    rows = []
    
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                cells.append(cell_text)
        if cells:
            rows.append(' | '.join(cells))
    
    return '\n'.join(rows)
